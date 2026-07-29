from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
FIG = ROOT / "outputs" / "figures"
OUT = ROOT / "output" / "docx" / "现代密码学课程报告_曾嘉祺_最终提交版.docx"


GREEK = {
    r"\alpha": "α",
    r"\rho": "ρ",
    r"\Delta": "Δ",
    r"\sigma": "σ",
    r"\lambda": "λ",
    r"\mu": "μ",
    r"\epsilon": "ε",
    r"\omega": "ω",
    r"\mathcal": "",
    r"\mathbb": "",
    r"\mathbb{R}": "R",
    r"\mathbb{Z}": "Z",
    r"\lfloor": "⌊",
    r"\rfloor": "⌋",
    r"\sum": "∑",
    r"\in": "∈",
    r"\leftarrow": "←",
    r"\rightarrow": "→",
    r"\mapsto": "↦",
    r"\ll": "≪",
    r"\times": "×",
    r"\bmod": "mod",
    r"\cdot": "·",
    r"\Pr": "Pr",
    r"\mathcal{N}": "N",
    r"\eta": "η",
    r"\pm": "±",
    r"\lceil": "⌈",
    r"\rceil": "⌉",
    r"\lVert": "‖",
    r"\rVert": "‖",
    r"\equiv": "≡",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\neq": "≠",
    r"\star": "⋆",
    r"\mid": " | ",
}


def extract_balanced_arg(s: str, start: int) -> tuple[str, int]:
    assert s[start] == "{"
    depth = 0
    out = []
    for i in range(start, len(s)):
        ch = s[i]
        if ch == "{":
            depth += 1
            if depth > 1:
                out.append(ch)
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return "".join(out), i + 1
            out.append(ch)
        else:
            out.append(ch)
    return "".join(out), len(s)


def strip_comments(tex: str) -> str:
    lines = []
    for line in tex.splitlines():
        if line.lstrip().startswith("%"):
            continue
        lines.append(re.sub(r"(?<!\\)%.*", "", line).rstrip())
    return "\n".join(lines)


def bib_order_and_refs(bbl: str) -> tuple[dict[str, int], list[str]]:
    bbl = re.sub(r"\\begin\{thebibliography\}\{[^}]*\}", "", bbl)
    bbl = re.sub(r"\\end\{thebibliography\}", "", bbl)
    items = re.split(r"\\bibitem\{([^}]+)\}", bbl)
    order: dict[str, int] = {}
    refs: list[str] = []
    idx = 1
    for i in range(1, len(items), 2):
        key = items[i]
        body = items[i + 1]
        order[key] = idx
        text = clean_latex(body)
        text = re.sub(r"\\newblock", " ", text)
        text = re.sub(r"\\em\s+([^{}]+)", r"\1", text)
        text = re.sub(r"\s+", " ", text).strip()
        refs.append(text)
        idx += 1
    return order, refs


def collect_labels(tex: str) -> dict[str, int]:
    labels: dict[str, int] = {}
    for env, prefix in [("figure", "fig"), ("table", "tab")]:
        n = 1
        for m in re.finditer(rf"\\begin\{{{env}\}}(.*?)\\end\{{{env}\}}", tex, re.S):
            label = re.search(r"\\label\{([^}]+)\}", m.group(1))
            if label:
                labels[label.group(1)] = n
            n += 1
    algo = re.search(r"\\begin\{algorithm\}.*?\\label\{([^}]+)\}", tex, re.S)
    if algo:
        labels[algo.group(1)] = 1
    return labels


def cite_repl(match: re.Match[str], bibnums: dict[str, int]) -> str:
    nums = []
    for key in match.group(1).split(","):
        key = key.strip()
        nums.append(str(bibnums.get(key, key)))
    return "[" + ",".join(nums) + "]"


def clean_latex(text: str, bibnums: dict[str, int] | None = None, labels: dict[str, int] | None = None) -> str:
    bibnums = bibnums or {}
    labels = labels or {}
    text = text.replace("\r", "")
    text = re.sub(r"\\cite\{([^}]+)\}", lambda m: cite_repl(m, bibnums), text)
    text = re.sub(r"\\ref\{([^}]+)\}", lambda m: str(labels.get(m.group(1), m.group(1))), text)
    text = re.sub(r"\\(zihao|kaishu|normalfont|bfseries|small|centering)\b(?:\{[^}]*\})?", "", text)
    text = text.replace(r"\%", "%").replace(r"\_", "_").replace(r"\&", "&")
    text = text.replace(r"\|", "‖").replace(r"\{", "{").replace(r"\}", "}")
    text = text.replace("~", " ").replace("--", "-")
    text = text.replace(r"\quad", " ").replace(r"\,", " ")
    for k, v in sorted(GREEK.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(k, v)
    text = re.sub(r"\\operatorname\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\widetilde\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathcal\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathbb\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\$([^$]+)\$", lambda m: clean_latex(m.group(1), bibnums, labels), text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{([^{}]*)\})?", lambda m: m.group(1) or "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def get_cmd_arg(tex: str, cmd: str) -> str:
    m = re.search(rf"\\{cmd}(?:\[[^\]]*\])?\s*\{{", tex)
    if not m:
        return ""
    arg, _ = extract_balanced_arg(tex, m.end() - 1)
    return clean_latex(arg)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_east_asia_font(target, font: str) -> None:
    if hasattr(target, "_element"):
        elem = target._element
    else:
        elem = target.element
    r_pr = elem.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.6)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    set_east_asia_font(normal, "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(0x1F, 0x4D, 0x78)),
        ("Heading 2", 12.5, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 3", 11.5, RGBColor(0x1F, 0x4D, 0x78)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        set_east_asia_font(style, "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_para(doc: Document, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    return p


def add_caption(doc: Document, text: str, keep_with_next: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(9)


def create_numbered_list(doc: Document) -> int:
    """Create a standalone decimal list so numbering restarts at one."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for name, value in [("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("lvlJc", "left")]:
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "420")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    p.add_run(text)
    return p


def extract_caption(block: str) -> str:
    m = re.search(r"\\caption(?:\[[^\]]*\])?\{", block)
    if not m:
        return ""
    arg, end = extract_balanced_arg(block, m.end() - 1)
    return clean_latex(arg)


def parse_tabular(block: str, bibnums: dict[str, int], labels: dict[str, int]) -> list[list[str]]:
    begin = block.find(r"\begin{tabular}")
    if begin < 0:
        return []
    pos = begin + len(r"\begin{tabular}")
    while pos < len(block) and block[pos].isspace():
        pos += 1
    if pos < len(block) and block[pos] == "{":
        _, pos = extract_balanced_arg(block, pos)
    end = block.find(r"\end{tabular}", pos)
    if end < 0:
        return []
    body = block[pos:end]
    rows: list[list[str]] = []
    for raw in re.split(r"\\\\", body):
        raw = raw.strip()
        raw = re.sub(r"\\(toprule|midrule|bottomrule|hline)", "", raw).strip()
        if not raw:
            continue
        raw = re.sub(r"\\multicolumn\{[^}]+\}\{[^}]+\}\{([^}]*)\}", r"\1", raw)
        cells = [clean_latex(c, bibnums, labels) for c in raw.split("&")]
        if cells and any(cells):
            rows.append(cells)
    return rows


def add_table_from_rows(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    for i, row in enumerate(rows):
        tr_pr = table.rows[i]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if i == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if j < len(row):
                cell.text = row[j]
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.keep_with_next = i < len(rows) - 1
                for r in p.runs:
                    r.font.size = Pt(8.5)
                    r.font.name = "Times New Roman"
                    set_east_asia_font(r, "宋体")
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def add_figure(doc: Document, block: str, fig_no: int) -> None:
    img = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", block)
    if img:
        path = (REPORT / img.group(1)).resolve()
        if not path.exists():
            path = (ROOT / img.group(1)).resolve()
        if path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(path), width=Inches(5.7))
    cap = extract_caption(block)
    if cap:
        add_caption(doc, f"图 {fig_no} {cap}")


def add_algorithm(doc: Document, block: str) -> None:
    cap = extract_caption(block) or "PP-SRSF 鲁棒安全聚合流程"
    add_caption(doc, f"算法 1 {cap}")
    lines = []
    for raw in block.splitlines():
        raw = raw.strip()
        if (
            not raw
            or raw == "}"
            or raw.startswith("\\begin")
            or raw.startswith("\\end")
            or raw.startswith("\\caption")
            or raw.startswith("\\label")
            or raw.startswith("\\zihao")
            or raw.startswith("\\LinesNumbered")
        ):
            continue
        raw = raw.replace(r"\;", "")
        if raw.startswith(r"\KwIn{"):
            arg, _ = extract_balanced_arg(raw, raw.find("{"))
            raw = "输入：" + arg
        elif raw.startswith(r"\KwOut{"):
            arg, _ = extract_balanced_arg(raw, raw.find("{"))
            raw = "输出：" + arg
        elif raw.startswith(r"\For{"):
            arg, _ = extract_balanced_arg(raw, raw.find("{"))
            raw = "对于" + arg + "，重复执行："
        raw = raw.rstrip("{}").strip()
        text = clean_latex(raw)
        if text:
            lines.append(text)
    num_id = create_numbered_list(doc)
    for line in lines:
        add_numbered_paragraph(doc, line, num_id)


def build_frontmatter(doc: Document, tex: str) -> None:
    title = get_cmd_arg(tex, "title")
    etitle = get_cmd_arg(tex, "etitle")
    author = get_cmd_arg(tex, "author")
    address = get_cmd_arg(tex, "address")
    abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", tex, re.S)
    keywords = re.search(r"\\begin\{keywords\}(.*?)\\end\{keywords\}", tex, re.S)
    eabstract = re.search(r"\\begin\{eabstract\}(.*?)\\end\{eabstract\}", tex, re.S)
    ekeywords = re.search(r"\\begin\{ekeywords\}(.*?)\\end\{ekeywords\}", tex, re.S)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    set_east_asia_font(r, "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(etitle)
    r.italic = True
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(author).bold = True
    p.add_run("\n" + address)

    if abstract:
        p = doc.add_paragraph()
        p.add_run("摘要：").bold = True
        p.add_run(clean_latex(abstract.group(1)))
    if keywords:
        p = doc.add_paragraph()
        p.add_run("关键词：").bold = True
        p.add_run(clean_latex(keywords.group(1)))
    if eabstract:
        p = doc.add_paragraph()
        p.add_run("Abstract: ").bold = True
        p.add_run(clean_latex(eabstract.group(1)))
    if ekeywords:
        p = doc.add_paragraph()
        p.add_run("Keywords: ").bold = True
        p.add_run(clean_latex(ekeywords.group(1)))


def split_body(tex: str) -> list[tuple[str, str]]:
    body = tex[tex.find(r"\section{引言}") : tex.find(r"\bibliographystyle")]
    pattern = re.compile(
        r"(\\begin\{figure\}.*?\\end\{figure\}|"
        r"\\begin\{table\}.*?\\end\{table\}|"
        r"\\begin\{algorithm\}.*?\\end\{algorithm\}|"
        r"\\begin\{proposition\}.*?\\end\{proposition\}|"
        r"\\begin\{proof\}.*?\\end\{proof\}|"
        r"\\begin\{enumerate\}.*?\\end\{enumerate\}|"
        r"\\\[.*?\\\])",
        re.S,
    )
    parts: list[tuple[str, str]] = []
    last = 0
    for m in pattern.finditer(body):
        if m.start() > last:
            parts.append(("text", body[last : m.start()]))
        token = m.group(0)
        if token.startswith(r"\begin{figure}"):
            kind = "figure"
        elif token.startswith(r"\begin{table}"):
            kind = "table"
        elif token.startswith(r"\begin{algorithm}"):
            kind = "algorithm"
        elif token.startswith(r"\begin{proposition}"):
            kind = "proposition"
        elif token.startswith(r"\begin{proof}"):
            kind = "proof"
        elif token.startswith(r"\begin{enumerate}"):
            kind = "enumerate"
        else:
            kind = "equation"
        parts.append((kind, token))
        last = m.end()
    if last < len(body):
        parts.append(("text", body[last:]))
    return parts


def add_text_block(doc: Document, block: str, bibnums: dict[str, int], labels: dict[str, int], counters: dict[str, int]) -> None:
    lines = [x.strip() for x in block.splitlines()]
    para = []
    for line in lines:
        if not line:
            if para:
                add_text_para(doc, " ".join(para), bibnums, labels)
                para = []
            continue
        sec = re.match(r"\\section\{([^}]+)\}", line)
        sub = re.match(r"\\subsection\{([^}]+)\}", line)
        subsub = re.match(r"\\subsubsection\{([^}]+)\}", line)
        if sec or sub or subsub:
            if para:
                add_text_para(doc, " ".join(para), bibnums, labels)
                para = []
            if sec:
                counters["section"] += 1
                counters["subsection"] = 0
                title = f"{counters['section']} {clean_latex(sec.group(1))}"
                doc.add_heading(title, level=1)
            elif sub:
                counters["subsection"] += 1
                title = f"{counters['section']}.{counters['subsection']} {clean_latex(sub.group(1))}"
                doc.add_heading(title, level=2)
            elif subsub:
                title = clean_latex(subsub.group(1))
                doc.add_heading(title, level=3)
        else:
            if line.startswith("\\") and not line.startswith("\\item"):
                continue
            para.append(line)
    if para:
        add_text_para(doc, " ".join(para), bibnums, labels)


def add_text_para(doc: Document, text: str, bibnums: dict[str, int], labels: dict[str, int]) -> None:
    text = clean_latex(text, bibnums, labels)
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)


def add_references(doc: Document, refs: list[str]) -> None:
    doc.add_heading("参考文献", level=1)
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.add_run(f"[{i}] {ref}")


def main() -> None:
    tex = strip_comments((REPORT / "main.tex").read_text(encoding="utf-8"))
    bbl = (REPORT / "main.bbl").read_text(encoding="utf-8")
    bibnums, refs = bib_order_and_refs(bbl)
    labels = collect_labels(tex)

    doc = Document()
    doc.core_properties.author = "曾嘉祺"
    doc.core_properties.title = "基于秘密共享的密态异常梯度过滤与鲁棒安全聚合机制研究"
    doc.core_properties.subject = "现代密码学课程报告"
    style_doc(doc)
    build_frontmatter(doc, tex)

    counters = {"section": 0, "subsection": 0}
    fig_no = 1
    tab_no = 1
    for kind, block in split_body(tex):
        if kind == "text":
            add_text_block(doc, block, bibnums, labels, counters)
        elif kind == "figure":
            add_figure(doc, block, fig_no)
            fig_no += 1
        elif kind == "table":
            cap = extract_caption(block)
            if cap:
                add_caption(doc, f"表 {tab_no} {cap}", keep_with_next=True)
            rows = parse_tabular(block, bibnums, labels)
            add_table_from_rows(doc, rows)
            tab_no += 1
        elif kind == "algorithm":
            add_algorithm(doc, block)
        elif kind == "equation":
            eq = block.replace(r"\[", "").replace(r"\]", "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(clean_latex(eq, bibnums, labels))
            r.font.name = "Cambria Math"
        elif kind == "proposition":
            text = re.sub(r"\\begin\{proposition\}|\\end\{proposition\}", "", block).strip()
            p = doc.add_paragraph()
            p.add_run("命题：").bold = True
            p.add_run(clean_latex(text, bibnums, labels))
        elif kind == "proof":
            text = re.sub(r"\\begin\{proof\}|\\end\{proof\}", "", block).strip()
            p = doc.add_paragraph()
            p.add_run("证明：").bold = True
            p.add_run(clean_latex(text, bibnums, labels))
        elif kind == "enumerate":
            body = re.sub(r"\\begin\{enumerate\}(?:\[[^]]*\])?|\\end\{enumerate\}", "", block)
            num_id = create_numbered_list(doc)
            for item in re.split(r"\\item", body):
                item = item.strip()
                if not item:
                    continue
                add_numbered_paragraph(doc, clean_latex(item, bibnums, labels), num_id)

    add_references(doc, refs)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("现代密码学课程报告")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
